from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm
import latex2mathml.converter
from lxml import etree

from bibliography_parser import book_text
from logger import log
from string import ascii_lowercase

from list_helper import list_number

def int_to_roman(num):
        val = [
            1000, 900, 500, 400,
            100, 90, 50, 40,
            10, 9, 5, 4,
            1
            ]
        syb = [
            "M", "CM", "D", "CD",
            "C", "XC", "L", "XL",
            "X", "IX", "V", "IV",
            "I"
            ]
        roman_num = ''
        i = 0
        while num > 0:
            for _ in range(num // val[i]):
                roman_num += syb[i]
                num -= val[i]
            i += 1
        return roman_num

class DocxGenerator:
    def __init__(self):
        self.document = Document('template.docx')
        self.last_line_empty = True
        self.reset_counters()
        self.structural_headings = []
        self.structural_headings_casefolded = []
        self.books = []
        self.tags = {}
        self.book_tags = {}
        self.table_numbering = 'global'
        self.picture_numbering = 'global'

        self.current_list = None
        
        xslt = etree.parse('mml2omml.xsl')
        self.math_transform = etree.XSLT(xslt)
        
        styles = self.document.styles
        paragraph_styles = [
            s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH
        ]
        for style in paragraph_styles:
            print(style.name)
    
    def reset_counters(self):
        self.current_section = 0
        self.current_subsections = [0, 0, 0, 0, 0]
        self.table_global_counter = 0
        self.picture_global_counter = 0
        self.table_section_counter = 0
        self.picture_section_counter = 0
        self.book_counter = 0

    def section_numbering(self, number):
        return str(self.current_section) + '.' + str(number)
    def section_numbering_full(self, level):
        return str(self.current_section) + '.' + '.'.join([str(x) for x in self.current_subsections[:level-1]])
    def picture_number(self, plus_one=False):
        add = 1 if plus_one else 0
        if self.picture_numbering == 'section':
            return self.section_numbering(self.picture_section_counter + add)
        else:
            return str(self.picture_global_counter + add)
    def table_number(self, plus_one=False):
        add = 1 if plus_one else 0
        if self.table_numbering == 'section':
            return self.section_numbering(self.table_section_counter + add)
        else:
            return str(self.table_global_counter + add)
    
    def set_settings(self, settings):
        self.settings = settings
        self.structural_headings = settings.get('structNames', [])
        self.structural_headings_casefolded = [s.casefold() for s in self.structural_headings]

    def generate_toc(self):
        self.generate_structural_heading(
            self.settings.get('tocTitle', 'Содержание'),
            is_toc=True)

        paragraph = self.document.add_paragraph()
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')  # creates a new element
        fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'   # change 1-3 depending on heading levels you need

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:t')
        fldChar3.text = "НАЖМИТЕ СЮДА ПРАВОЙ КНОПКОЙ И ОБНОВИТЕ"
        fldChar2.append(fldChar3)

        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')

        r_element = run._r
        r_element.append(fldChar)
        r_element.append(instrText)
        r_element.append(fldChar2)
        r_element.append(fldChar4)
        p_element = paragraph._p

    def generate_ast(self, ast):
        self.register_ast(ast)
        self.reset_counters()

        for x in ast:
            self.generate_object(x)
    
    def generate_object(self, obj):
        if obj['type'] == 'heading':
            self.generate_heading(obj)
        elif obj['type'] == 'paragraph':
            self.generate_paragraph(obj)
        elif obj['type'] == 'list':
            self.generate_list(obj)
        elif obj['type'] == 'block_equation':
            self.generate_block_equation(obj)
        elif obj['type'] == 'block_table':
            self.generate_block_table(obj)
        elif obj['type'] == 'picture':
            self.generate_picture(obj)
        elif obj['type'] == 'toc':
            self.generate_toc()
        elif obj['type'] == 'block_code':
            self.generate_block_code(obj)
        elif obj['type'] == 'book':
            self.generate_book(obj)
        elif obj['type'] == 'biblio':
            self.generate_bibliography()
        #elif obj['type'] == 'book_data':
        #    self.book_tags = obj['tags']
    
    def pad_if_needed(self, key):
        padding = self.settings.get(key, True)
        if padding and not self.last_line_empty:
            self.document.add_paragraph('')
            self.last_line_empty = True

    def generate_structural_heading(self, text, is_toc=False):
        should_break_page = self.settings.get('sectionPageBreak', True)
        if should_break_page:
            self.document.add_page_break()
            self.last_line_empty = True

        self.pad_if_needed('headingsArePadded')
        heading = self.document.add_heading(text, level=1)
        if is_toc:
            heading.style = self.document.styles['TOC Heading']
        self.last_line_empty = text == ""
        self.pad_if_needed('headingsArePadded')

    def generate_heading(self, obj):
        text = self.aggregate_text(obj['children'])
        level = obj['level']
        if level == 1 and text.casefold() in self.structural_headings_casefolded:
            actual_text = [
                self.structural_headings[i] for i in range(len(self.structural_headings))
                if text.casefold() == self.structural_headings_casefolded[i]
                ][0]
            return self.generate_structural_heading(actual_text)

        
        should_autonumber = self.settings.get('headingsAutonumber', True)
        should_break_page = self.settings.get('sectionPageBreak', True)

        if level == 1:
            self.current_section += 1
            self.current_subsections = [0, 0, 0, 0, 0]
            self.table_section_counter = 0
            self.picture_section_counter = 0

            if should_break_page:
                self.document.add_page_break()
                self.last_line_empty = True
            if should_autonumber and text[0] not in '123456789':
                text = str(self.current_section) + ' ' + text
        else:
            self.current_subsections[level-2] += 1
            for i in range(level-1, len(self.current_subsections)):
                self.current_subsections[i] = 0
            if should_autonumber and text[0] not in '123456789':
                text = self.section_numbering_full(level) + ' ' + text

        self.pad_if_needed('headingsArePadded')
        heading = self.document.add_heading(text, level=level)
        if level == 1:
            heading.style = self.document.styles['Section header']
        self.last_line_empty = text == ""
        self.pad_if_needed('headingsArePadded')
        
    def generate_paragraph(self, obj):
        text = self.aggregate_text(obj['children'])
        paragraph = self.document.add_paragraph()
        self.add_text(paragraph, obj['children'])
        self.last_line_empty = text == ""

    def list_number(self, number, level):
        if level == 1:
            return str(number)
        elif level == 2:
            return ascii_lowercase[number - 1]
        elif level == 3:
            return int_to_roman(number).lower()
        else:
            return str(number)

    def generate_list(self, obj):
        ordered = obj['ordered']
        level = obj['level']
        list_style = self.settings.get('listStyle', 'word')
        if list_style == 'text':
            for i, elem in enumerate(obj['children'], start=1):
                text = self.aggregate_text(elem['children'])
                prefix = (level-1) * '\t'
                if ordered:
                    prefix += self.list_number(i, level) + ') '
                else:
                    prefix += '- '
                if text[-1] != ',' and i != len(obj['children']):
                    text += ','
                #if paragraph is None:
                paragraph = self.document.add_paragraph()
                paragraph.add_run(prefix)
                self.add_text(paragraph, elem['children'])
        elif list_style == 'word':
            if ordered:
                style = 'Main Number List'
            else:
                style = 'Bullet List'
            if level > 1:
                style += ' ' + str(level)

            for elem in obj['children']:
                paragraph = self.document.add_paragraph(style=style)
                list_number(self.document, paragraph, self.current_list, level - 1, ordered)
                paragraph.style = 'List Paragraph'
                self.current_list = paragraph
                self.add_text(paragraph, elem['children'])
            if level == 1:
                self.current_list = None

        self.last_line_empty = False
    def generate_equation(self, run, obj):
        placeholder = self.settings.get('equationPlaceholder', 'ВВЕДИТЕ УРАВНЕНИЕ')
        if obj['children'] == '':
            text = placeholder
            mathml = None
        else:
            text = obj['children']
            if text.startswith('!'):
                text = placeholder + '(' + text[1:] + ')'
                mathml = None
            else:
                mathml = latex2mathml.converter.convert(text)
                tree = etree.fromstring(mathml)
                new_tree = self.math_transform(tree)
                run._r.append(new_tree.getroot())
                return
        mainTextElem = OxmlElement('m:t')
        mainTextElem.text = text
        runElem = OxmlElement('m:r')
        runElem.append(mainTextElem)
        mathElem = OxmlElement('m:oMath')
        mathElem.append(runElem)
        run._r.append(mathElem)
    def generate_block_equation(self, obj):
        self.pad_if_needed('blockEquationsArePadded')
        mainTextElem = OxmlElement('m:t')
        if 'text' in obj:
            text = obj['text']
        else:
            text = self.aggregate_text(obj['children'])
        if text == '':
            mainTextElem.text = self.settings.get('equationPlaceholder', 'ВВЕДИТЕ УРАВНЕНИЕ')
        else:
            mathml = latex2mathml.converter.convert(text)
            tree = etree.fromstring(mathml)
            new_tree = self.math_transform(tree)
            paragraph = self.document.add_paragraph()
            paragraph._p.append(new_tree.getroot())
            return
        runElem = OxmlElement('m:r')
        runElem.append(mainTextElem)
        mathElem = OxmlElement('m:oMath')
        mathElem.append(runElem)
        mathParaElem = OxmlElement('m:oMathPara')
        mathParaElem.append(mathElem)
        paragraph = self.document.add_paragraph()
        paragraph._p.append(mathParaElem)
        self.last_line_empty = False
        self.pad_if_needed('blockEquationsArePadded')
    def generate_block_table(self, obj):
        self.pad_if_needed('tablesArePadded')
        self.table_global_counter += 1
        self.table_section_counter += 1
        table_text = self.settings.get('tableName', 'Таблица')
        table_text += ' ' + self.table_number() + ' – ' + obj['text']
        self.document.add_paragraph(table_text)
        table = self.document.add_table(cols=obj['width'], rows=obj['height'])
        table.style = 'Table Grid'
        self.last_line_empty = False
        self.pad_if_needed('tablesArePadded')
    
    def generate_picture(self, obj):
        self.pad_if_needed('picturesArePadded')

        width = Cm(obj['size']) if obj['symbol'] == 'w' else None
        height = Cm(obj['size']) if obj['symbol'] == 'h' else None

        self.picture_global_counter += 1
        self.picture_section_counter += 1
        picture_text = self.settings.get('pictureName', 'Рисунок') 
        picture_text += ' ' + self.picture_number() + ' – '
        path = obj['path']
        try:
            self.document.add_picture(path, width=width, height=height)
        except FileNotFoundError:
            self.document.add_paragraph("File " + path + " not found!")
            log("WARNING: File " + path + " not found!")
        para = self.document.paragraphs[-1]
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        para = self.document.add_paragraph(picture_text)
        self.add_text(para, obj['children'])
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.last_line_empty = False
        self.pad_if_needed('picturesArePadded')

    def generate_block_code(self, obj):
        paragraph = self.document.add_paragraph(obj['text'].rstrip())
        paragraph.style = self.document.styles['Code']
        self.last_line_empty = False

    def generate_book(self, obj):
        self.books.append(book_text(obj))

    def aggregate_text(self, objs):
        result = ''
        for obj in objs:
            if (obj['type'] == 'equation' or obj['type'] == 'book_ref' or obj['type'] == 'ref' 
                    or obj['type'] == 'inline_code'):
                result += obj['children']
                continue
            if 'text' in obj:
                result += obj['text']
            if 'children' in obj:
                result += self.aggregate_text(obj['children'])
        return result
    def add_text(self, paragraph, objs, strong=False, emphasis=False):
        for obj in objs:
            if obj['type'] == 'equation':
                run = paragraph.add_run()
                run.bold = strong
                run.italic = emphasis
                self.generate_equation(run, obj)
                continue
            elif obj['type'] == 'book_ref':
                tag = obj['children']
                index = "???"
                if tag in self.book_tags:
                    index = self.book_tags[tag]
                run = paragraph.add_run(f"[{index}]")
                continue
            elif obj['type'] == 'ref':
                tag = obj['children']
                index = "???"
                if tag == 't':
                    index = self.table_number(plus_one=True)
                elif tag == 'p':
                    index = self.picture_number(plus_one=True)
                elif tag in self.tags:
                    index = self.tags[tag]
                run = paragraph.add_run(index)
                continue
            elif obj['type'] == 'codespan':
                run = paragraph.add_run(obj['text'])
                run.bold = strong
                run.italic = emphasis
                run.style = 'InlineCode'
                continue
            elif obj['type'] == 'list':
                self.generate_list(obj)
                continue
            if 'text' in obj:
                if obj['text'].lstrip()[:2] == "$$" and obj['text'].rstrip()[-2:] == "$$":
                    obj['text'] = obj['text'][2:-2]
                    self.generate_block_equation(obj)
                    continue
                run = paragraph.add_run(obj['text'])
                run.bold = strong
                run.italic = emphasis
            if 'children' in obj:
                next_strong = obj['type'] == 'strong'
                next_emphasis = obj['type'] == 'emphasis'
                self.add_text(paragraph, obj['children'], 
                    strong=next_strong, emphasis=next_emphasis)
    
    def generate_bibliography(self):
        self.generate_structural_heading(
            self.settings.get('biblioTitle', 'СПИСОК ИСПОЛЬЗУЕМЫХ ИСТОЧНИКОВ'))
        for i, book in enumerate(self.books, start=1):
            self.document.add_paragraph(str(i) + '. ' + book)

    def register_ast(self, ast):
        for x in ast:
            self.register_object(x)

    def register_object(self, obj):
        if obj['type'] == 'heading':
            if obj['level'] == 1:
                text = self.aggregate_text(obj['children'])
                if text.casefold() in self.structural_headings_casefolded:
                    return
                self.current_section += 1
                self.current_subsections = [0, 0, 0, 0, 0]
                self.table_section_counter = 0
                self.picture_section_counter = 0
            else:
                self.current_subsections[obj['level']-2] += 1
                for i in range(obj['level']-1, len(self.current_subsections)):
                    self.current_subsections[i] = 0
        elif obj['type'] == 'block_table':
            self.table_global_counter += 1
            self.table_section_counter += 1
            if obj['tag']:
                if obj['tag'] in self.tags:
                    log('ERROR: tag ' + obj['tag'] + 'is used twice!')
                else:
                    self.tags[obj['tag']] = self.table_number()
        elif obj['type'] == 'picture':
            self.picture_global_counter += 1
            self.picture_section_counter += 1
            if obj['tag']:
                if obj['tag'] in self.tags:
                    log('ERROR: tag ' + obj['tag'] + 'is used twice!')
                else:
                    self.tags[obj['tag']] = self.picture_number()
        elif obj['type'] == 'book':
            self.book_counter += 1
            if obj['tag']:
                if obj['tag'] in self.book_tags:
                    log('ERROR: book tag ' + obj['tag'] + 'is used twice!')
                else:
                    self.book_tags[obj['tag']] = str(self.book_counter)
        elif obj['type'] == 'table_numbering':
            self.table_numbering = self.aggregate_text(obj['children'])
        elif obj['type'] == 'picture_numbering':
            self.picture_numbering = self.aggregate_text(obj['children'])

    def save(self, filename):
        try:
            self.document.save(filename)
        except:
            log("ERROR: Couldn't save file! (maybe it's currently opened in some editor?)")